# Portal use to extract keywords from the CIE textbook references from C1 to C20, start develop on 25th Aug 2026 *************
import datetime
import io
import os
import fitz  # PyMuPDF
import streamlit as st

# Word Document Libraries
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Google API Libraries
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

# ==========================================
# 0. STREAMLIT PAGE CONFIG & CUSTOM STYLING
# ==========================================
st.set_page_config(
    page_title="9618 Computer Science Portal", 
    page_icon="💻",
    layout="wide"
)

# Custom Styling Theme
st.markdown("""
    <style>
    .stApp { background-color: #E6BBFC !important; }
    [data-testid="stSidebar"] { background-color: #C663F8 !important; }
    [data-testid="stSidebar"] > div:first-child { background-color: #C663F8 !important; }
    div[data-baseweb="input"], div[data-baseweb="select"] > div, .stTextInput input, .stSelectbox select {
        background-color: #BBFCBC !important;
        color: #070F9C !important;
        border-radius: 8px !important;
        border: 1px solid #620092 !important;
    }
    label, .stWidgetLabel p { color: #2D004B !important; font-weight: bold !important; }
    input { color: #070F9C !important; }
    </style>
""", unsafe_allow_html=True)


# ==========================================
# 1. CONFIGURATION & DIRECTORY SETUP
# ==========================================
SYLLABUS_CODE = "9618"
SCOPES = ['https://www.googleapis.com/auth/drive.readonly']
BASE_DIR = "9618HandOuts"

# Folder paths mapped under the main 9618HandOuts directory
LOCAL_FOLDERS = {
    "paper1": os.path.join(BASE_DIR, "9618P1C1_C8"),
    "paper2": os.path.join(BASE_DIR, "9618P2C9_C12"),
    "paper3": os.path.join(BASE_DIR, "9618P3C13_C16"),
    "paper4": os.path.join(BASE_DIR, "9618P4C17_C20"),
    "other_notes": os.path.join(BASE_DIR, "Other9618Notes")
}

# Auto-filter Chapter breakdown for syllabus 9618
PAPER_CHAPTER_MAPPING = {
    "paper1": [f"Chapter {i}" for i in range(1, 9)],    # Chapters 1 to 8
    "paper2": [f"Chapter {i}" for i in range(9, 13)],   # Chapters 9 to 12
    "paper3": [f"Chapter {i}" for i in range(13, 17)],  # Chapters 13 to 16
    "paper4": [f"Chapter {i}" for i in range(17, 21)]   # Chapters 17 to 20
}

# Create local directories automatically if they do not exist
for folder_path in LOCAL_FOLDERS.values():
    os.makedirs(folder_path, exist_ok=True)


# ==========================================
# 2. GOOGLE DRIVE SERVICE & SYNC ENGINE
# ==========================================
def build_drive_service():
    """Authenticates using Google Service Account credentials stored in st.secrets."""
    try:
        if "gcp_service_account" in st.secrets:
            service_account_info = dict(st.secrets["gcp_service_account"])
            creds = service_account.Credentials.from_service_account_info(
                service_account_info, 
                scopes=SCOPES
            )
            return build('drive', 'v3', credentials=creds)
        else:
            st.error("❌ Missing [gcp_service_account] section in secrets.toml.")
            return None
    except Exception as e:
        st.error(f"❌ Service Account Authentication Error: {e}")
        return None

def sync_drive_folder_to_local(folder_key: str) -> tuple[int, str]:
    """Downloads missing PDF files from Google Drive to the target local folder."""
    service = build_drive_service()
    if not service:
        return 0, "Failed to authenticate Service Account."
    
    folder_ids = st.secrets.get("drive_folders", {})
    drive_folder_id = folder_ids.get(folder_key)
    
    if not drive_folder_id:
        return 0, f"Missing drive_folder_id for `{folder_key}` in secrets.toml."

    local_path = LOCAL_FOLDERS[folder_key]
    
    try:
        query = f"'{drive_folder_id}' in parents and trashed = false"
        drive_files = []
        page_token = None

        while True:
            response = service.files().list(
                q=query,
                fields="nextPageToken, files(id, name, mimeType)",
                pageToken=page_token,
                pageSize=100
            ).execute()
            
            drive_files.extend(response.get('files', []))
            page_token = response.get('nextPageToken', None)
            if not page_token:
                break

        downloaded_count = 0
        for file_info in drive_files:
            file_name = file_info['name']
            file_id = file_info['id']
            local_file_path = os.path.join(local_path, file_name)

            if not os.path.exists(local_file_path):
                request = service.files().get_media(fileId=file_id)
                with open(local_file_path, "wb") as f:
                    downloader = MediaIoBaseDownload(f, request)
                    done = False
                    while not done:
                        _, done = downloader.next_chunk()
                downloaded_count += 1

        total_local = len([f for f in os.listdir(local_path) if os.path.isfile(os.path.join(local_path, f))])
        return downloaded_count, f"Synced {downloaded_count} new file(s) for `{folder_key}` (Total local: {total_local})."
        
    except Exception as e:
        return 0, f"Sync error for `{folder_key}`: {e}"

def perform_bulk_sync():
    """Syncs all 5 local directories with their corresponding Google Drive folders."""
    total_synced = 0
    messages = []
    for f_key in LOCAL_FOLDERS.keys():
        count, msg = sync_drive_folder_to_local(f_key)
        total_synced += count
        messages.append(msg)
    return total_synced, messages


# ==========================================
# 3. HELPER FUNCTIONS (DOCX & PDF PREVIEW)
# ==========================================

def add_page_number_to_run(run):
    """Inserts native Word field codes for dynamic page numbers in headers."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def create_worksheet_docx(basket_items: list) -> io.BytesIO:
    """Generates the dynamic Word document (.docx) handout with corrected image scaling."""
    doc = Document()
    section = doc.sections[0]

    # Page Margins & Setup (Standard Letter / A4 setup)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Dynamic Top Header
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_p.add_run("Page ")
    add_page_number_to_run(header_run)

    # Add Main Title Document Heading
    main_heading = doc.add_heading(f'PTES {SYLLABUS_CODE} Computer Science Handout', level=1)
    main_heading.paragraph_format.space_after = Inches(0.1)

    for idx, item in enumerate(basket_items):
        # Add heading for individual page source
        heading = doc.add_heading(f"Source: {item['file']} (Page {item['page'] + 1})", level=2)
        heading.paragraph_format.space_before = Inches(0.1)
        heading.paragraph_format.space_after = Inches(0.1)

        # Open and render PDF page
        pdf_doc = fitz.open(item['path'])
        page = pdf_doc.load_page(item['page'])
        
        # Render high-resolution pixmap image
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img_data = io.BytesIO(pix.tobytes("png"))

        # Scaled image width to 6.5 inches to ensure heading + image fit on 1 page
        img_paragraph = doc.add_paragraph()
        img_paragraph.paragraph_format.space_after = Inches(0.0)
        run = img_paragraph.add_run()
        run.add_picture(img_data, width=Inches(6.5))

        # Add page break only between items (not after the final item)
        if idx < len(basket_items) - 1:
            doc.add_page_break()
            
        pdf_doc.close()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
#==============================================================================================================================

def render_pdf_page_preview(filepath: str, page_num: int = 0):
    """Renders a single PDF page into PNG byte format for display."""
    try:
        doc = fitz.open(filepath)
        page = doc.load_page(page_num)
        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
        img_bytes = pix.tobytes("png")
        doc.close()
        return img_bytes
    except Exception as e:
        st.error(f"Unable to render page preview: {e}")
        return None

def execute_chapter_search(paper_key: str, keyword_string: str, selected_chapter: str) -> list[dict]:
    """
    Performs full-text search across BOTH the specified Paper folder 
    AND the shared Other9618Notes folder.
    """
    results = []
    keywords = [k.strip().lower() for k in keyword_string.split(",") if k.strip()]
    
    # Folders to scan: Paper Specific Directory + Other Notes Directory
    target_folders = [LOCAL_FOLDERS[paper_key], LOCAL_FOLDERS["other_notes"]]

    for folder_path in target_folders:
        if os.path.exists(folder_path):
            for file in os.listdir(folder_path):
                if file.endswith(".pdf"):
                    # Apply Chapter Filtering if a specific chapter is selected
                    if selected_chapter != "All Chapters":
                        chap_num = selected_chapter.split(" ")[1]  # Extract number e.g. '1'
                        file_lower = file.lower()
                        # Checks matching patterns: ch1, chapter1, chapter 1
                        if f"ch{chap_num}" not in file_lower and f"chapter{chap_num}" not in file_lower and f"chapter {chap_num}" not in file_lower:
                            # Skip files that don't match the chapter filter (unless in Other Notes)
                            if folder_path != LOCAL_FOLDERS["other_notes"]:
                                continue

                    filepath = os.path.join(folder_path, file)
                    try:
                        doc = fitz.open(filepath)
                        for page_num in range(len(doc)):
                            text = doc[page_num].get_text().lower()
                            if all(kw in text for kw in keywords):
                                results.append({
                                    "file": file, 
                                    "page": page_num, 
                                    "path": filepath
                                })
                        doc.close()
                    except Exception:
                        continue
    return results


# ==========================================
# 4. APP STATE INITIALIZATION
# ==========================================
if 'handout_basket' not in st.session_state:
    st.session_state.handout_basket = []

for p in ["paper1", "paper2", "paper3", "paper4"]:
    if f"{p}_results" not in st.session_state:
        st.session_state[f"{p}_results"] = []

if 'has_auto_synced' not in st.session_state:
    st.session_state.has_auto_synced = True
    with st.spinner("🚀 Waking up portal & auto-syncing files from Google Drive..."):
        perform_bulk_sync()


# ==========================================
# 5. STREAMLIT UI LAYOUT & TABS
# ==========================================
st.title("PUSAT TINGKATAN ENAM SENGKURONG")
st.subheader(f"💻 {SYLLABUS_CODE} Computer Science Topical Portal")

# Sidebar Controls
with st.sidebar:
    st.title("⚙️ Control Panel")
    st.markdown("---")
    
    st.subheader("🔄 Google Drive Sync")
    if st.button("🔄 Sync Google Drive", type="primary", use_container_width=True):
        with st.spinner("Syncing Google Drive folders..."):
            count, msgs = perform_bulk_sync()
            st.success(f"🎉 Sync Complete! {count} new file(s) downloaded.")
            for m in msgs:
                st.caption(m)

    st.markdown("---")
    st.subheader("📊 Basket Summary")
    st.metric(label="Saved Pages in Basket", value=len(st.session_state.handout_basket))

    if st.button("🗑️ Clear Entire Basket", use_container_width=True):
        st.session_state.handout_basket = []
        st.toast("Basket cleared successfully.")
        st.rerun()

# 6 Navigation Tabs
tabs = st.tabs([
    "📘 Theory P1 topics", 
    "📗 Theory P2 topics", 
    "📙 Theory P3 topics", 
    "📕 Practical P4", 
    "🛒 Notes/Cart", 
    "⚙️ Upload/Admin"
])

def render_paper_tab(tab_object, paper_key: str, paper_title: str):
    """Generates UI elements for Paper Tabs 1 to 4."""
    with tab_object:
        st.header(f"🔍 Search {paper_title}")
        
        col_kw, col_ch = st.columns([2, 1])
        with col_kw:
            kw = st.text_input(f"Enter Keyword", placeholder="e.g., binary, stack, subroutines", key=f"{paper_key}_kw")
        with col_ch:
            chapter_options = ["All Chapters"] + PAPER_CHAPTER_MAPPING[paper_key]
            selected_ch = st.selectbox("Select Chapter Filter", options=chapter_options, key=f"{paper_key}_chap")

        if st.button(f"Search {paper_title}", type="primary", key=f"btn_{paper_key}"):
            if kw.strip():
                with st.spinner(f"Scanning {paper_title} & Other Notes..."):
                    st.session_state[f"{paper_key}_results"] = execute_chapter_search(paper_key, kw, selected_ch)
            else:
                st.warning("Please enter a keyword.")

        results = st.session_state[f"{paper_key}_results"]
        if results:
            st.write(f"Found **{len(results)}** matching page(s):")
            for idx, item in enumerate(results):
                with st.expander(f"📄 {item['file']} | Page {item['page'] + 1}"):
                    c1, c2 = st.columns([3, 1])
                    with c1:
                        preview_img = render_pdf_page_preview(item["path"], item["page"])
                        if preview_img:
                            st.image(preview_img, use_container_width=True)
                    with c2:
                        if st.button("➕ Add to Cart", key=f"add_{paper_key}_{idx}"):
                            st.session_state.handout_basket.append(item)
                            st.toast(f"Added Page {item['page'] + 1} to basket!")
                            st.rerun()
                        
                        st.markdown("<br>", unsafe_allow_html=True)
                        with open(item["path"], "rb") as pdf_f:
                            st.download_button(
                                label="📥 Download Target PDF",
                                data=pdf_f,
                                file_name=item["file"],
                                mime="application/pdf",
                                key=f"dl_{paper_key}_{idx}"
                            )

# Render Paper Tabs
render_paper_tab(tabs[0], "paper1", "Paper 1 (Chapters 1–8)")
render_paper_tab(tabs[1], "paper2", "Paper 2 (Chapters 9–12)")
render_paper_tab(tabs[2], "paper3", "Paper 3 (Chapters 13–16)")
render_paper_tab(tabs[3], "paper4", "Paper 4 (Chapters 17–20)")


# --- TAB 5: BASKET / CART ---
with tabs[4]:
    st.header("🛒 Worksheet Basket & Handout Generator")
    
    if st.session_state.handout_basket:
        st.subheader("Selected Pages in Cart")
        for idx, item in enumerate(st.session_state.handout_basket):
            c_info, c_action = st.columns([4, 1])
            with c_info:
                st.markdown(f"📄 **Item {idx + 1}:** `{item['file']}` — **Page {item['page'] + 1}**")
            with c_action:
                if st.button("🔴 DELETE", key=f"del_cart_{idx}"):
                    st.session_state.handout_basket.pop(idx)
                    st.toast(f"Removed item {idx + 1} from basket.")
                    st.rerun()
            st.markdown("---")

        st.subheader("📝 Export Options")
        doc_buffer = create_worksheet_docx(st.session_state.handout_basket)
        st.download_button(
            label="🪄 Download Dynamic Word Document Handout",
            data=doc_buffer,
            file_name=f"{SYLLABUS_CODE}_Topical_Handout.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True
        )
    else:
        st.info("🛒 Your cart is currently empty. Search for topics across Paper Tabs 1–4 and click '➕ Add to Cart'.")


# --- TAB 6: UPLOAD & ADMIN DASHBOARD ---
with tabs[5]:
    st.header("⚙️ Upload & Admin Dashboard")
    st.caption("Access Google Drive upload destinations for adding new chapter PDFs, notes, and references.")

    admin_pwd = st.secrets.get("ADMIN_PASSWORD", "")
    pwd_input = st.text_input("Enter Admin Password", type="password")

    if pwd_input and pwd_input == admin_pwd:
        st.success("🔓 Authenticated as Administrator")
        st.markdown("---")
        drive_links = st.secrets.get("drive_web_links", {})

        col_a, col_b = st.columns(2)
        with col_a:
            st.link_button("📘 Open Paper 1 Drive (Ch 1–8)", drive_links.get("paper1", "https://drive.google.com"), type="primary", use_container_width=True)
            st.markdown("<br>", unsafe_allow_html=True)
            st.link_button("📙 Open Paper 3 Drive (Ch 13–16)", drive_links.get("paper3", "https://drive.google.com"), type="primary", use_container_width=True)
            st.markdown("<br>", unsafe_allow_html=True)
            st.link_button("📁 Open Other Notes Drive Folder", drive_links.get("other_notes", "https://drive.google.com"), type="primary", use_container_width=True)
        with col_b:
            st.link_button("📗 Open Paper 2 Drive (Ch 9–12)", drive_links.get("paper2", "https://drive.google.com"), type="primary", use_container_width=True)
            st.markdown("<br>", unsafe_allow_html=True)
            st.link_button("📕 Open Paper 4 Drive (Ch 17–20)", drive_links.get("paper4", "https://drive.google.com"), type="primary", use_container_width=True)
    elif pwd_input:
        st.error("❌ Incorrect Admin Password.")


# ==========================================
# 6. PORTAL FOOTER
# ==========================================
st.markdown("---")
SCHOOL_NAME = "Pusat Tingkatan Enam Sengkurong (PTES)"
SCHOOL_VISION = "Nurturing Resilient Leaders & Future-Ready Citizens"

footer_html = f"""
<div style="text-align: center; padding: 15px 0px; color: #2D004B; font-family: sans-serif;">
    <p style="margin: 0; font-size: 1.0em; font-weight: bold;">🏫 {SCHOOL_NAME}</p>
    <p style="margin: 5px 0; font-size: 0.9em; font-style: italic;">"{SCHOOL_VISION}"</p>
</div>
"""
st.markdown(footer_html, unsafe_allow_html=True)
