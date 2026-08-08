# ********** Computer Science 9618 PYP Portal ***********
import io
import os
import fitz  # PyMuPDF
import streamlit as st

# Word Document Libraries
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Google API Libraries
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

# ==========================================
# 0. STREAMLIT PAGE CONFIG & CUSTOM STYLING
# ==========================================
st.set_page_config(
    page_title="9618 Computer Science PYP Archives", 
    page_icon="💻",
    layout="wide"
)

# Custom CSS Theme Implementation
st.markdown("""
    <style>
    .stApp {
        background-color: #E6BBFC !important;
    }
    [data-testid="stSidebar"] {
        background-color: #C663F8 !important;
    }
    [data-testid="stSidebar"] > div:first-child {
        background-color: #C663F8 !important;
    }
    div[data-baseweb="input"], 
    div[data-baseweb="select"] > div, 
    .stTextInput input, 
    .stSelectbox select {
        background-color: #BBFCBC !important;
        color: #070F9C !important;
        border-radius: 8px !important;
        border: 1px solid #620092 !important;
    }
    label, .stWidgetLabel p {
        color: #2D004B !important;
        font-weight: bold !important;
    }
    input {
        color: #070F9C !important;
    }
    </style>
""", unsafe_allow_html=True)


# ==========================================
# 1. CONFIGURATION & DIRECTORY SETUP
# ==========================================
SYLLABUS_CODE = "9618"
SCOPES = ['https://www.googleapis.com/auth/drive.readonly']

LOCAL_FOLDERS = {
    "theory": "9618Theory",
    "practical": "9618Practical",
    "answer_scheme": "9618AnswerScheme",
    "zips": "9618Zip"
}

for folder_path in LOCAL_FOLDERS.values():
    os.makedirs(folder_path, exist_ok=True)


# ==========================================
# 2. SERVICE ACCOUNT AUTHENTICATION & SYNC
# ==========================================
def build_drive_service():
    """Authenticates using Google Service Account credentials stored in st.secrets."""
    try:
        if "gcp_service_account" in st.secrets:
            service_account_info = dict(st.secrets["gcp_service_account"])
            creds = service_account.Credentials.from_service_account_info(
                service_account_info, 
                scopes=SCOPES
            )
            return build('drive', 'v3', credentials=creds)
        else:
            st.error("❌ Missing [gcp_service_account] section in secrets.toml.")
            return None
    except Exception as e:
        st.error(f"❌ Service Account Authentication Error: {e}")
        return None

def sync_drive_folder_to_local(folder_key: str) -> tuple[int, str]:
    """Downloads missing files from Google Drive to local directories."""
    service = build_drive_service()
    if not service:
        return 0, "Failed to authenticate Service Account."
    
    folder_ids = st.secrets.get("drive_folders", {})
    drive_folder_id = folder_ids.get(folder_key)
    
    if not drive_folder_id:
        return 0, f"Missing drive_folder_id for `{folder_key}` in secrets.toml."

    local_path = LOCAL_FOLDERS[folder_key]
    try:
        query = f"'{drive_folder_id}' in parents and trashed = false"
        results = service.files().list(q=query, fields="files(id, name, mimeType)").execute()
        drive_files = results.get('files', [])
        downloaded_count = 0

        for file_info in drive_files:
            file_name = file_info['name']
            file_id = file_info['id']
            local_file_path = os.path.join(local_path, file_name)

            if not os.path.exists(local_file_path):
                request = service.files().get_media(fileId=file_id)
                with open(local_file_path, "wb") as f:
                    downloader = MediaIoBaseDownload(f, request)
                    done = False
                    while not done:
                        _, done = downloader.next_chunk()
                downloaded_count += 1

        return downloaded_count, f"Synced {downloaded_count} file(s) for `{folder_key}`."
    except Exception as e:
        return 0, f"Sync error for `{folder_key}`: {e}"

def perform_bulk_sync():
    """Syncs all 4 configured Google Drive folders."""
    total_synced = 0
    messages = []
    for f_key in ["theory", "practical", "answer_scheme", "zips"]:
        count, msg = sync_drive_folder_to_local(f_key)
        total_synced += count
        messages.append(msg)
    return total_synced, messages


# ==========================================
# 3. HELPER FUNCTIONS
# ==========================================
def add_page_number_to_run(run):
    """Adds a dynamic word field for Page numbers in header/footer."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def render_pdf_page_preview(filepath: str, page_num: int = 0):
    """Renders a single PDF page into PNG bytes for preview."""
    try:
        doc = fitz.open(filepath)
        page = doc.load_page(page_num)
        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
        img_bytes = pix.tobytes("png")
        doc.close()
        return img_bytes
    except Exception as e:
        st.error(f"Unable to render page preview: {e}")
        return None

def execute_pdf_search(folder_key: str, keyword_string: str) -> list[dict]:
    """Performs full-text search across a specific local folder and returns page matches."""
    results = []
    keywords = [k.strip().lower() for k in keyword_string.split(",") if k.strip()]
    folder_path = LOCAL_FOLDERS[folder_key]
    
    if os.path.exists(folder_path):
        for file in os.listdir(folder_path):
            if file.endswith(".pdf"):
                filepath = os.path.join(folder_path, file)
                try:
                    doc = fitz.open(filepath)
                    for page_num in range(len(doc)):
                        text = doc[page_num].get_text().lower()
                        if all(kw in text for kw in keywords):
                            results.append({
                                "file": file, 
                                "page": page_num, 
                                "path": filepath
                            })
                    doc.close()
                except Exception:
                    continue
    return results


# ==========================================
# 4. APP STATE INITIALIZATION
# ==========================================
if 'handout_basket' not in st.session_state:
    st.session_state.handout_basket = []

if 'theory_search_results' not in st.session_state:
    st.session_state.theory_search_results = []
if 'practical_search_results' not in st.session_state:
    st.session_state.practical_search_results = []

if 'has_auto_synced' not in st.session_state:
    st.session_state.has_auto_synced = True
    with st.spinner("🚀 Waking up portal & auto-syncing files via Service Account..."):
        perform_bulk_sync()


# ==========================================
# 5. STREAMLIT UI LAYOUT
# ==========================================
st.title("PUSAT TINGKATAN ENAM SENGKURONG")
st.subheader(f"💻 {SYLLABUS_CODE} Computer Science PYP Resource Library")

# --- SIDEBAR CONTROLS ---
with st.sidebar:
    st.header("🔄 Google Drive Sync")
    if st.button("🔄 Sync Google Drive", type="primary", use_container_width=True):
        with st.spinner("Syncing Google Drive folders..."):
            synced_count, sync_msgs = perform_bulk_sync()
            st.success(f"🎉 Sync Complete! {synced_count} new file(s) downloaded.")
            for m in sync_msgs:
                st.caption(m)

    st.markdown("---")
    st.metric(label="Saved Pages in Basket", value=len(st.session_state.handout_basket))

    if st.button("🗑️ Clear Entire Basket", use_container_width=True):
        st.session_state.handout_basket = []
        st.rerun()

# --- RE-ORDERED NAVIGATION TABS ---
tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "1. 🔍 Search Keyword Theory", 
    "2. 💻 Search Keyword Practical", 
    "3. 🛒 HandOut / Cart", 
    "4. 📦 Source File", 
    "5. 🔑 Answer Scheme", 
    "6. ⚙️ Upload PYP Admin"
])


# --- TAB 1: SEARCH KEYWORD THEORY ---
with tab1:
    st.header("🔍 Search Keyword (Theory - Papers 1 & 3)")
    t_kw = st.text_input("Enter Keywords for Theory", placeholder="e.g., binary tree, recursion, pipeline", key="theory_kw")

    if st.button("Search Theory Papers", type="primary", key="btn_search_theory"):
        if t_kw.strip():
            with st.spinner("Scanning Theory PDFs..."):
                st.session_state.theory_search_results = execute_pdf_search("theory", t_kw)
        else:
            st.warning("Please enter a keyword.")

    if st.session_state.theory_search_results:
        st.write(f"Found **{len(st.session_state.theory_search_results)}** matching page(s):")
        for idx, item in enumerate(st.session_state.theory_search_results):
            with st.expander(f"📄 {item['file']} | Page {item['page'] + 1}"):
                c1, c2 = st.columns([3, 1])
                with c1:
                    preview_img = render_pdf_page_preview(item["path"], item["page"])
                    if preview_img:
                        st.image(preview_img, use_container_width=True)
                with c2:
                    if st.button("➕ Add to Basket", key=f"add_t_{idx}"):
                        st.session_state.handout_basket.append(item)
                        st.toast(f"Added Page {item['page'] + 1} to basket!")
                        st.rerun()
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    with open(item["path"], "rb") as pdf_f:
                        st.download_button(
                            label="📥 Download Full PDF",
                            data=pdf_f,
                            file_name=item["file"],
                            mime="application/pdf",
                            key=f"dl_t_{idx}"
                        )


# --- TAB 2: SEARCH KEYWORD PRACTICAL ---
with tab2:
    st.header("💻 Search Keyword (Practical - Paper 4)")
    p_kw = st.text_input("Enter Keywords for Practical", placeholder="e.g., OOP, stacks, file handling", key="practical_kw")

    if st.button("Search Practical Papers", type="primary", key="btn_search_practical"):
        if p_kw.strip():
            with st.spinner("Scanning Practical PDFs..."):
                st.session_state.practical_search_results = execute_pdf_search("practical", p_kw)
        else:
            st.warning("Please enter a keyword.")

    if st.session_state.practical_search_results:
        st.write(f"Found **{len(st.session_state.practical_search_results)}** matching page(s):")
        for idx, item in enumerate(st.session_state.practical_search_results):
            with st.expander(f"📄 {item['file']} | Page {item['page'] + 1}"):
                c1, c2 = st.columns([3, 1])
                with c1:
                    preview_img = render_pdf_page_preview(item["path"], item["page"])
                    if preview_img:
                        st.image(preview_img, use_container_width=True)
                with c2:
                    if st.button("➕ Add to Basket", key=f"add_p_{idx}"):
                        st.session_state.handout_basket.append(item)
                        st.toast(f"Added Page {item['page'] + 1} to basket!")
                        st.rerun()
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    with open(item["path"], "rb") as pdf_f:
                        st.download_button(
                            label="📥 Download Full PDF",
                            data=pdf_f,
                            file_name=item["file"],
                            mime="application/pdf",
                            key=f"dl_p_{idx}"
                        )


# --- TAB 3: HANDOUT / CART ---
with tab5 if False else tab3:
    st.header("🛒 HandOut / Cart Management")
    
    if st.session_state.handout_basket:
        st.subheader("Selected Pages in Your Cart")
        st.markdown("Review your items below. Click **DELETE** to remove an individual page.")
        
        # Display each item with an individual Delete button
        for idx, item in enumerate(st.session_state.handout_basket):
            col_info, col_action = st.columns([4, 1])
            with col_info:
                st.markdown(f"📄 **Item {idx + 1}:** `{item['file']}` — **Page {item['page'] + 1}**")
            with col_action:
                if st.button("🔴 DELETE", key=f"del_item_{idx}"):
                    st.session_state.handout_basket.pop(idx)
                    st.toast(f"Removed item {idx + 1} from cart.")
                    st.rerun()
            st.markdown("---")

        st.markdown("<br>", unsafe_allow_html=True)
        st.subheader("📝 Export Options")
        
        if st.button("🪄 Export Handout to Word (.docx)", type="primary", use_container_width=True):
            try:
                doc = Document()
                section = doc.sections[0]
                section.orientation = WD_ORIENT.PORTRAIT
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)

                header = section.header
                header_p = header.paragraphs[0]
                header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                header_run = header_p.add_run("Page ")
                add_page_number_to_run(header_run)

                doc.add_heading(f'PTES {SYLLABUS_CODE} Computer Science Worksheet', level=1)

                for idx, item in enumerate(st.session_state.handout_basket):
                    doc.add_heading(f"Source: {item['file']} (Page {item['page'] + 1})", level=2)
                    pdf_doc = fitz.open(item['path'])
                    page = pdf_doc.load_page(item['page'])
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img_data = io.BytesIO(pix.tobytes("png"))
                    doc.add_picture(img_data, width=Inches(6.5))
                    if idx < len(st.session_state.handout_basket) - 1:
                        doc.add_page_break()
                    pdf_doc.close()

                target_filename = f"{SYLLABUS_CODE}_CS_Worksheet.docx"
                doc.save(target_filename)
                with open(target_filename, "rb") as f:
                    st.download_button(
                        label="📥 Download Generated Word Document", 
                        data=f, 
                        file_name=target_filename,
                        type="primary",
                        use_container_width=True
                    )
            except Exception as e:
                st.error(f"Error generating Word file: {e}")
    else:
        st.info("🛒 Your cart is currently empty. Search for questions in Tab 1 or Tab 2 and click '➕ Add to Basket' to add pages here.")


# --- TAB 4: SOURCE FILE ---
with tab4:
    st.header("📦 Download Practical Source Files & Evidence")
    st.caption("Select the exam parameters to locate and download practical ZIP archives.")

    sf_col1, sf_col2, sf_col3 = st.columns(3)
    with sf_col1:
        sf_year = st.selectbox("Select Year", [str(y) for y in range(2026, 2020, -1)], key="sf_yr")
    with sf_col2:
        sf_month = st.selectbox("Select Session", ["June (s)", "November (w)"], key="sf_mth")
        sf_m_code = "s" if "June" in sf_month else "w"
    with sf_col3:
        sf_variant = st.selectbox(
            "Select Practical Variant", 
            ["41", "42", "43"], 
            index=2 if "June" in sf_month else 1,
            key="sf_var"
        )

    sf_short_year = sf_year[-2:]
    
    possible_zip_names = [
        f"{SYLLABUS_CODE}_{sf_m_code}{sf_short_year}_sf_{sf_variant}.zip",
        f"{SYLLABUS_CODE}_{sf_m_code}{sf_short_year}_zip_{sf_variant}.zip",
        f"{SYLLABUS_CODE}_{sf_m_code}{sf_short_year}_source_{sf_variant}.zip"
    ]

    st.markdown("---")
    found_zip_path = None
    matched_zip_name = ""

    for zip_candidate in possible_zip_names:
        candidate_path = os.path.join(LOCAL_FOLDERS["zips"], zip_candidate)
        if os.path.exists(candidate_path):
            found_zip_path = candidate_path
            matched_zip_name = zip_candidate
            break

    if found_zip_path:
        st.success(f"✅ Found Practical Source File Archive: `{matched_zip_name}`")
        with open(found_zip_path, "rb") as zip_f:
            st.download_button(
                label=f"📥 Download Source File Archive ({matched_zip_name})",
                data=zip_f,
                file_name=matched_zip_name,
                mime="application/zip",
                type="primary"
            )
    else:
        st.warning(f"⚠️ Source file archive for Year {sf_year}, Session {sf_month}, Variant {sf_variant} was not found in local storage.")


# --- TAB 5: ANSWER SCHEME ---
with tab5:
    st.header("🔑 Answer Scheme Finder (Marking Schemes)")
    col_y, col_m, col_v = st.columns(3)
    with col_y:
        as_year = st.selectbox("Select Year", [str(y) for y in range(2026, 2020, -1)], key="as_yr")
    with col_m:
        as_month = st.selectbox("Select Session", ["June (s)", "November (w)"], key="as_mth")
        month_code = "s" if "June" in as_month else "w"
    with col_v:
        as_variant = st.selectbox("Select Variant Component", ["11", "12", "13", "21", "22", "23", "31", "32", "33", "41", "42", "43"], key="as_var")

    short_year = as_year[-2:]
    expected_ms_filename = f"{SYLLABUS_CODE}_{month_code}{short_year}_ms_{as_variant}.pdf"

    st.markdown("---")
    found_ms_path = None
    for folder_path in [LOCAL_FOLDERS["answer_scheme"], LOCAL_FOLDERS["theory"], LOCAL_FOLDERS["practical"]]:
        check_path = os.path.join(folder_path, expected_ms_filename)
        if os.path.exists(check_path):
            found_ms_path = check_path
            break

    if found_ms_path:
        st.success(f"✅ Found Answer Scheme: `{expected_ms_filename}`")
        with open(found_ms_path, "rb") as f:
            st.download_button("📥 Download Answer Scheme PDF", f, file_name=expected_ms_filename, mime="application/pdf", type="primary")
        
        with st.expander("👁️ Preview Answer Scheme Document"):
            doc = fitz.open(found_ms_path)
            for p in range(len(doc)):
                img_data = render_pdf_page_preview(found_ms_path, p)
                if img_data:
                    st.image(img_data, caption=f"Page {p + 1}", use_container_width=True)
            doc.close()
    else:
        st.warning(f"⚠️ Answer Scheme `{expected_ms_filename}` was not found locally.")


# --- TAB 6: UPLOAD PYP ADMIN ---
with tab6:
    st.header("⚙️ Upload PYP Admin Dashboard")
    st.caption("Direct shortcuts to Google Drive folder dashboards for fast file uploads and management.")

    admin_pwd = st.secrets.get("ADMIN_PASSWORD", "")
    pwd_input = st.text_input("Enter Admin Password", type="password")

    if pwd_input and pwd_input == admin_pwd:
        st.success("🔓 Authenticated as Administrator")
        st.markdown("---")
        st.subheader("📁 Google Drive Upload Dashboards")

        drive_links = st.secrets.get("drive_web_links", {})

        col_a, col_b = st.columns(2)
        with col_a:
            st.link_button("📘 Open Theory Drive Folder", drive_links.get("theory", "https://drive.google.com"), type="primary", use_container_width=True)
            st.markdown("<br>", unsafe_allow_html=True)
            st.link_button("🔑 Open Answer Scheme Drive Folder", drive_links.get("answer_scheme", "https://drive.google.com"), type="primary", use_container_width=True)

        with col_b:
            st.link_button("💻 Open Practical Drive Folder", drive_links.get("practical", "https://drive.google.com"), type="primary", use_container_width=True)
            st.markdown("<br>", unsafe_allow_html=True)
            st.link_button("📦 Open Source File Zip Drive Folder", drive_links.get("zips", "https://drive.google.com"), type="primary", use_container_width=True)

    elif pwd_input:
        st.error("❌ Incorrect Admin Password.")

# ==========================================
# 6. PORTAL FOOTER
# ==========================================
st.markdown("---")
SCHOOL_NAME = "Pusat Tingkatan Enam Sengkurong (PTES)"
SCHOOL_VISION = "Nurturing Resilient Leaders & Future-Ready Citizens"
DEVELOPER_NAME = "Miss Hajah Nurul Haziqah HN / Computer Science Department"

footer_html = f"""
<div style="text-align: center; padding: 15px 0px; color: #2D004B; font-family: sans-serif;">
    <p style="margin: 0; font-size: 1.0em; font-weight: bold;">🏫 {SCHOOL_NAME}</p>
    <p style="margin: 5px 0; font-size: 0.9em; font-style: italic;">"{SCHOOL_VISION}"</p>
    <p style="margin: 5px 0 0 0; font-size: 0.85em; font-weight: 600; color: #070F9C;">💻 Developed by {DEVELOPER_NAME}</p>
</div>
"""
st.markdown(footer_html, unsafe_allow_html=True)
